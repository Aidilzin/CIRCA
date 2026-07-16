"""
utils/convert_markdown_thesis_to_docx.py
=========================================
Reusable utility module to compile formatted Word Documents (.docx) from Markdown
sources using python-docx. Applies UiTM thesis standards (TNR font, custom paragraph
spacings, page numbers, justified alignment, and styled table cells).

Can be imported as a Python module or executed directly from the command line.
"""

import argparse
import logging
import os
import re
import sys
import time
from typing import Dict, Any, Optional

import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Explicit imports for type hints to bypass docx attribute resolution errors
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Cell

logger = logging.getLogger("circa.utils.markdown_docx")


def set_cell_background(cell: _Cell, fill_color: str) -> None:
    """Sets the background color of a table cell (e.g., 'F2F2F2')."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)


def set_cell_margins(cell: _Cell, top: int = 100, bottom: int = 100, left: int = 150, right: int = 150) -> None:
    """Sets internal padding (margins) for a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table: Table, color: str = "D3D3D3", sz: str = "4", val: str = "single") -> None:
    """Applies clean horizontal borders to a Word table according to formatting standards."""
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)


def add_page_number(run: Run) -> None:
    """Adds a page number field (PAGE) inside a footer run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)


def parse_inline_formatting(paragraph: Paragraph, text: str) -> None:
    """Parses standard markdown bold (**), italic (*), and code (`) and adds formatted runs to paragraph."""
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|https?://\S+)')
    tokens = pattern.split(text)
    
    for token in tokens:
        if not token:
            continue
        
        run = paragraph.add_run()
        run.font.name = 'Times New Roman'
        
        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(199, 37, 78) # Pink/red code color
        elif token.startswith('http://') or token.startswith('https://'):
            run.text = token
            run.font.color.rgb = RGBColor(0, 0, 238) # Link blue
            run.underline = True
        else:
            run.text = token
            run.font.size = Pt(12)


def convert_markdown_to_docx(input_filepath: str, output_filepath: str) -> Dict[str, Any]:
    """
    Reads a Markdown file and compiles it into a cleanly styled Word (.docx) document
    using Academic formatting standards (Times New Roman, 1.5 line spacing, custom margin
    setup, and dynamically styled tables).

    Args:
        input_filepath: Path to the source Markdown document (.md).
        output_filepath: Output file destination path (.docx).

    Returns:
        Dict[str, Any] containing conversion status, metadata, and error details if failed.
    """
    start_time = time.time()
    result = {
        "success": False,
        "output_path": output_filepath,
        "error": None,
        "metadata": {
            "paragraphs_count": 0,
            "tables_count": 0,
            "duration_seconds": 0.0,
            "lines_processed": 0
        }
    }

    if not os.path.exists(input_filepath):
        result["error"] = f"Source file does not exist: {input_filepath}"
        return result

    try:
        with open(input_filepath, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        num_lines = len(lines)
        doc = docx.Document()
        
        # 1. Page Margin Setup based on UiTM formatting guidelines
        # Top margin: 3 cm, Bottom: 2 cm, Left: 4 cm, Right: 2.5 cm
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(4)
            section.right_margin = Cm(2.5)
            
            # Center-aligned page numbering in the footer
            footer = section.footer
            p_foot = footer.paragraphs[0]
            p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_foot.paragraph_format.space_before = Pt(0)
            p_foot.paragraph_format.space_after = Pt(0)
            run_foot = p_foot.add_run()
            run_foot.font.name = 'Times New Roman'
            run_foot.font.size = Pt(12)
            add_page_number(run_foot)
            
        # Set default style properties (Times New Roman, 12pt, justified)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0x00, 0x00, 0x00) # True Black for thesis print
        
        # Create custom caption styles for TOC mapping
        styles = doc.styles
        if "Figure Caption" not in styles:
            fig_style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
            fig_style.base_style = styles["Normal"]
            fig_style.font.name = 'Times New Roman'
            fig_style.font.size = Pt(11)
            fig_style.font.italic = True
            fig_style.font.color.rgb = RGBColor(0, 0, 0)
            
        if "Table Caption" not in styles:
            tab_style = styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
            tab_style.base_style = styles["Normal"]
            tab_style.font.name = 'Times New Roman'
            tab_style.font.size = Pt(12)
            tab_style.font.bold = True
            tab_style.font.color.rgb = RGBColor(0, 0, 0)

        # Process blocks
        i = 0
        while i < num_lines:
            line = lines[i]
            line_stripped = line.strip()
            
            # 1. Headings
            if line_stripped.startswith('#'):
                h_match = re.match(r'^(#+)\s*(.*)$', line_stripped)
                if h_match:
                    level = len(h_match.group(1))
                    title = h_match.group(2)
                    
                    if level == 1:
                        # CHAPTER TITLE
                        parts = title.split(':', 1)
                        
                        # Insert Page Break before Chapter Title if there's prior content
                        if doc.paragraphs and len(doc.paragraphs) > 1:
                            p_break = doc.add_paragraph()
                            p_break.paragraph_format.space_before = Pt(0)
                            p_break.paragraph_format.space_after = Pt(0)
                            p_break.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
                            
                        p_chap = doc.add_paragraph()
                        p_chap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_chap.paragraph_format.space_before = Pt(18)
                        p_chap.paragraph_format.space_after = Pt(12)
                        p_chap.paragraph_format.keep_with_next = True
                        
                        run1 = p_chap.add_run(parts[0].strip().upper())
                        run1.font.name = 'Times New Roman'
                        run1.font.size = Pt(14)
                        run1.bold = True
                        
                        if len(parts) > 1:
                            p_title = doc.add_paragraph()
                            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_title.paragraph_format.space_before = Pt(0)
                            p_title.paragraph_format.space_after = Pt(24)
                            p_title.paragraph_format.keep_with_next = True
                            
                            run2 = p_title.add_run(parts[1].strip().upper())
                            run2.font.name = 'Times New Roman'
                            run2.font.size = Pt(14)
                            run2.bold = True
                            
                    else:
                        # Section/Subsection Title
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_before = Pt(18)
                        p.paragraph_format.space_after = Pt(6)
                        p.paragraph_format.keep_with_next = True
                        
                        run = p.add_run(title)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                        run.bold = True
                        
                i += 1
                continue
                
            # 2. Blockquotes / Alerts
            if line_stripped.startswith('>'):
                quote_lines = []
                while i < num_lines and lines[i].strip().startswith('>'):
                    quote_line = lines[i].strip()[1:].strip()
                    quote_line = re.sub(r'^\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', '', quote_line).strip()
                    if quote_line:
                        quote_lines.append(quote_line)
                    i += 1
                    
                quote_text = " ".join(quote_lines)
                if quote_text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    
                    run = p.add_run("Note: ")
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                    
                    parse_inline_formatting(p, quote_text)
                    for r in p.runs[1:]:
                        r.italic = True
                        r.font.size = Pt(11)
                        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                continue
                
            # 3. Horizontal Rules
            if line_stripped in ('***', '---', '___'):
                i += 1
                continue
                
            # 4. Images
            img_match = re.search(r'!\[(.*?)\]\((.*?)\)', line_stripped)
            if img_match:
                caption = img_match.group(1)
                img_rel_path = img_match.group(2)
                
                md_dir = os.path.dirname(input_filepath)
                img_abs_path = os.path.normpath(os.path.join(md_dir, img_rel_path))
                
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(4)
                p_img.paragraph_format.keep_with_next = True
                
                if os.path.exists(img_abs_path):
                    try:
                        run_img = p_img.add_run()
                        run_img.add_picture(img_abs_path, width=Inches(5.5))
                        logger.info("Embedded image: %s", img_abs_path)
                    except Exception as e:
                        run_err = p_img.add_run(f"[Error loading image: {img_rel_path} - {str(e)}]")
                        run_err.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    run_err = p_img.add_run(f"[Image Placeholder: {img_rel_path}]")
                    run_err.italic = True
                    run_err.font.color.rgb = RGBColor(128, 128, 128)
                    logger.warning("Image file not found: %s", img_abs_path)
                    
                p_cap = doc.add_paragraph(style='Figure Caption')
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(12)
                run_cap = p_cap.add_run(caption)
                run_cap.italic = True
                run_cap.font.name = 'Times New Roman'
                run_cap.font.size = Pt(11)
                run_cap.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                
                i += 1
                continue
                
            # 5. Tables
            if line_stripped.startswith('|'):
                table_lines = []
                while i < num_lines and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                    
                if len(table_lines) >= 2:
                    parsed_rows = []
                    for t_line in table_lines:
                        cells = [c.strip() for c in t_line.split('|')]
                        if cells and cells[0] == '':
                            cells = cells[1:]
                        if cells and cells[-1] == '':
                            cells = cells[:-1]
                        parsed_rows.append(cells)
                        
                    header_row = parsed_rows[0]
                    data_rows = parsed_rows[1:]
                    if data_rows and all(re.match(r'^:?-+:?$', c) for c in data_rows[0]):
                        data_rows = data_rows[1:]
                        
                    num_cols = len(header_row)
                    num_rows_total = 1 + len(data_rows)
                    
                    table = doc.add_table(rows=num_rows_total, cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.allow_autofit = False
                    table.autofit = False
                    set_table_borders(table)
                    
                    col_widths = []
                    font_sz = 10
                    cell_margin_left_right = 150  # dxa
                    cell_margin_top_bottom = 100  # dxa
                    
                    first_hdr_lower = header_row[0].lower()
                    if num_cols == 6 and "id" in first_hdr_lower:
                        col_widths = [0.4, 1.4, 1.1, 0.9, 0.8, 1.12]
                    elif num_cols == 4 and "split" in first_hdr_lower:
                        col_widths = [1.42, 1.4, 1.5, 1.4]
                    elif num_cols == 4 and "metric" in first_hdr_lower and "verdict" in header_row[3].lower():
                        col_widths = [1.5, 1.2, 1.0, 2.02]
                    elif num_cols == 4 and "metric" in first_hdr_lower:
                        col_widths = [2.42, 1.1, 1.1, 1.1]
                    elif num_cols == 4 and "preprocessing stage" in first_hdr_lower:
                        col_widths = [1.5, 2.42, 0.9, 0.9]
                    elif num_cols == 14:
                        col_widths = [0.35, 0.4, 0.4, 0.4, 0.4, 0.4, 0.47, 0.45, 0.4, 0.35, 0.35, 0.35, 0.35, 0.35]
                        font_sz = 7.5
                        cell_margin_left_right = 60
                        cell_margin_top_bottom = 40
                    elif num_cols == 5 and "parameter" in first_hdr_lower:
                        col_widths = [1.1, 0.8, 0.8, 0.92, 2.1]
                    elif num_cols == 7 and "model variant" in first_hdr_lower:
                        col_widths = [1.4, 0.7, 0.7, 0.7, 0.7, 0.7, 0.82]
                    elif num_cols == 6 and "partition" in first_hdr_lower:
                        col_widths = [1.0, 1.0, 1.3, 0.8, 0.8, 0.82]
                    elif num_cols == 8 and "model variant" in first_hdr_lower:
                        col_widths = [1.1, 0.6, 0.6, 0.6, 0.6, 0.7, 0.7, 0.82]
                    elif num_cols == 7 and "model variant" in first_hdr_lower and "absolute" in header_row[2].lower():
                        col_widths = [0.9, 0.6, 0.7, 0.6, 0.6, 0.7, 1.62]
                    elif num_cols == 7 and "device / runtime" in header_row[2].lower():
                        col_widths = [0.9, 0.6, 1.4, 0.7, 0.7, 0.7, 0.72]
                    elif num_cols == 7 and "runtime device" in header_row[1].lower():
                        col_widths = [0.9, 0.9, 0.7, 0.8, 0.8, 0.8, 0.82]
                    elif num_cols == 10 and "model variant" in first_hdr_lower:
                        col_widths = [0.75, 0.45, 0.5, 0.5, 0.55, 0.55, 0.55, 0.55, 0.55, 0.77]
                        font_sz = 8.5
                    elif num_cols == 6 and "defect class" in first_hdr_lower:
                        col_widths = [1.0, 1.32, 0.85, 0.85, 0.85, 0.85]
                    elif num_cols == 4 and "defect class" in first_hdr_lower:
                        col_widths = [1.2, 1.2, 1.2, 2.12]
                    elif num_cols == 7 and "system" in first_hdr_lower:
                        col_widths = [0.9, 0.9, 1.0, 0.9, 0.5, 0.5, 1.02]
                    elif num_cols == 5 and "author and year" in first_hdr_lower:
                        col_widths = [1.1, 1.1, 1.1, 1.1, 1.32]
                    elif num_cols == 3 and "technique" in first_hdr_lower:
                        col_widths = [1.7, 1.5, 2.52]
                    else:
                        col_width = 5.72 / num_cols
                        col_widths = [col_width] * num_cols
                    
                    # Format Header Row
                    hdr_cells = table.rows[0].cells
                    for col_idx, col_text in enumerate(header_row):
                        cell = hdr_cells[col_idx]
                        set_cell_background(cell, "F2F2F2")
                        set_cell_margins(cell, top=cell_margin_top_bottom, bottom=cell_margin_top_bottom, left=cell_margin_left_right, right=cell_margin_left_right)
                        
                        p = cell.paragraphs[0]
                        p.text = ""
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing = 1.0
                        
                        run = p.add_run(col_text)
                        run.bold = True
                        run.font.size = Pt(font_sz + 0.5)
                        run.font.name = 'Times New Roman'
                        
                    # Format Data Rows
                    for r_idx, row_data in enumerate(data_rows):
                        row_cells = table.rows[r_idx + 1].cells
                        for col_idx in range(num_cols):
                            cell_val = row_data[col_idx] if col_idx < len(row_data) else ""
                            cell = row_cells[col_idx]
                            set_cell_margins(cell, top=cell_margin_top_bottom, bottom=cell_margin_top_bottom, left=cell_margin_left_right, right=cell_margin_left_right)
                            
                            if r_idx % 2 == 1:
                                set_cell_background(cell, "FAFAFA")
                                
                            p = cell.paragraphs[0]
                            p.text = ""
                            p.paragraph_format.space_before = Pt(2)
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.line_spacing = 1.0
                            
                            # Alignment Heuristics
                            stripped_val = cell_val.strip()
                            if re.match(r'^\d+(\.\d+)?%?$', stripped_val) or stripped_val.replace('+', '').replace('-', '').replace('pp', '').replace('×', '').strip().replace('.', '').isdigit():
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif stripped_val in ('-', 'No', 'Yes', 'COMPLETE', 'PENDING', 'Baseline_Vanilla', 'Baseline_CIRCA', 'HPO_7class'):
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                            parse_inline_formatting(p, cell_val)
                            
                            for run in p.runs:
                                run.font.size = Pt(font_sz)
                                run.font.name = 'Times New Roman'
                                
                    # Set Widths
                    for row in table.rows:
                        for col_idx, w in enumerate(col_widths):
                            if col_idx < len(row.cells):
                                row.cells[col_idx].width = Inches(w)
                                
                    p_after = doc.add_paragraph()
                    p_after.paragraph_format.space_after = Pt(6)
                continue
                
            # 6. Bullet Lists
            list_match = re.match(r'^(\s*)[-\*]\s+(.*)$', line)
            if list_match:
                indent_spaces = len(list_match.group(1))
                list_text = list_match.group(2)
                
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.5
                
                if indent_spaces >= 4:
                    p.paragraph_format.left_indent = Inches(0.5)
                elif indent_spaces >= 2:
                    p.paragraph_format.left_indent = Inches(0.25)
                    
                parse_inline_formatting(p, list_text)
                i += 1
                continue
                
            # 7. Standard Paragraphs
            if line_stripped:
                para_lines = [line_stripped]
                i += 1
                while i < num_lines:
                    next_line = lines[i].strip()
                    if (not next_line or 
                        next_line.startswith('#') or 
                        next_line.startswith('|') or 
                        next_line.startswith('>') or 
                        next_line.startswith('- ') or 
                        next_line.startswith('* ') or
                        next_line.startswith('![')):
                        break
                    para_lines.append(next_line)
                    i += 1
                    
                para_text = " ".join(para_lines)
                if para_text.startswith('**Table'):
                    p = doc.add_paragraph(style='Table Caption')
                else:
                    p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.5
                
                if para_text.startswith('[PENDING:') or para_text.startswith('[Insert Table') or para_text.startswith('[Insert Figure'):
                    remark = " [REMARK: This table/figure cannot be generated yet as the underlying experiment (Phases 4–7) has not been run. These results require training the final model weights first.]"
                    run = p.add_run(para_text + remark)
                    run.font.color.rgb = RGBColor(180, 50, 50)
                    run.bold = True
                    run.italic = True
                    run.font.size = Pt(11)
                elif para_text == '[INSERT LIST OF FIGURES HERE]':
                    p.text = ""
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5), WD_TABLE_ALIGNMENT.RIGHT)
                    run_hdr = p.add_run("Figure\tPage")
                    run_hdr.bold = True
                    run_hdr.font.name = 'Times New Roman'
                    run_hdr.font.size = Pt(12)
                    
                    p_field = doc.add_paragraph()
                    p_field.paragraph_format.space_before = Pt(6)
                    p_field.paragraph_format.space_after = Pt(12)
                    run = p_field.add_run()
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    fldChar1 = OxmlElement('w:fldChar')
                    fldChar1.set(qn('w:fldCharType'), 'begin')
                    instrText = OxmlElement('w:instrText')
                    instrText.set(qn('xml:space'), 'preserve')
                    instrText.text = 'TOC \\h \\z \\t "Figure Caption,1"'
                    fldChar2 = OxmlElement('w:fldChar')
                    fldChar2.set(qn('w:fldCharType'), 'separate')
                    fldChar3 = OxmlElement('w:fldChar')
                    fldChar3.set(qn('w:fldCharType'), 'end')
                    run._r.append(fldChar1)
                    run._r.append(instrText)
                    run._r.append(fldChar2)
                    run._r.append(fldChar3)
                elif para_text == '[INSERT LIST OF TABLES HERE]':
                    p.text = ""
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5), WD_TABLE_ALIGNMENT.RIGHT)
                    run_hdr = p.add_run("Table\tPage")
                    run_hdr.bold = True
                    run_hdr.font.name = 'Times New Roman'
                    run_hdr.font.size = Pt(12)
                    
                    p_field = doc.add_paragraph()
                    p_field.paragraph_format.space_before = Pt(6)
                    p_field.paragraph_format.space_after = Pt(12)
                    run = p_field.add_run()
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    fldChar1 = OxmlElement('w:fldChar')
                    fldChar1.set(qn('w:fldCharType'), 'begin')
                    instrText = OxmlElement('w:instrText')
                    instrText.set(qn('xml:space'), 'preserve')
                    instrText.text = 'TOC \\h \\z \\t "Table Caption,1"'
                    fldChar2 = OxmlElement('w:fldChar')
                    fldChar2.set(qn('w:fldCharType'), 'separate')
                    fldChar3 = OxmlElement('w:fldChar')
                    fldChar3.set(qn('w:fldCharType'), 'end')
                    run._r.append(fldChar1)
                    run._r.append(instrText)
                    run._r.append(fldChar2)
                    run._r.append(fldChar3)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    parse_inline_formatting(p, para_text)
                continue
                
            i += 1

        # Save document
        try:
            logger.info("Saving document to %s...", output_filepath)
            doc.save(output_filepath)
            result["success"] = True
        except PermissionError:
            logger.warning("Permission denied on %s, trying fallback...", output_filepath)
            fallback_path = output_filepath.replace(".docx", "_fallback.docx")
            doc.save(fallback_path)
            result["success"] = True
            result["output_path"] = fallback_path
            result["error"] = "Primary file locked. Saved to fallback path."
            
        result["metadata"]["paragraphs_count"] = len(doc.paragraphs)
        result["metadata"]["tables_count"] = len(doc.tables)
        result["metadata"]["lines_processed"] = num_lines
        
    except Exception as e:
        result["error"] = str(e)
        logger.error("Failed to convert markdown to docx", exc_info=True)

    result["metadata"]["duration_seconds"] = round(time.time() - start_time, 4)
    return result


if __name__ == '__main__':
    # Standalone CLI Guard
    parser = argparse.ArgumentParser(
        description="Convert Academic Markdown documents to styled Word (.docx) documents with template guidelines."
    )
    parser.add_argument(
        "-i", "--input", type=str, default="docs/CIRCA_THESIS_CH1-5.md",
        help="Path to source Markdown document (default: docs/CIRCA_THESIS_CH1-5.md)"
    )
    parser.add_argument(
        "-o", "--output", type=str, default="docs/2023276732_AIDIL_FYP THESIS.docx",
        help="Path to save output Word document (default: docs/2023276732_AIDIL_FYP THESIS.docx)"
    )
    
    args = parser.parse_args()
    
    # Configure root logging for CLI usage feedback
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
        handlers=[logging.StreamHandler(sys.stdout)]
    )
    
    print(f"[*] Starting conversion of '{args.input}' to '{args.output}'...")
    res = convert_markdown_to_docx(args.input, args.output)
    
    if res["success"]:
        print(f"[+] SUCCESS: Document compiled at: {res['output_path']}")
        print(f"    - Processed lines: {res['metadata']['lines_processed']}")
        print(f"    - Total paragraphs: {res['metadata']['paragraphs_count']}")
        print(f"    - Total tables: {res['metadata']['tables_count']}")
        print(f"    - Completed in {res['metadata']['duration_seconds']:.2f} seconds.")
        if res["error"]:
            print(f"    - Note: {res['error']}")
    else:
        print(f"[X] FAILED: {res['error']}", file=sys.stderr)
        sys.exit(1)
